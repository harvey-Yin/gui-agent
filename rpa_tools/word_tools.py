"""
Word文档处理工具
基于python-docx和docxtpl实现Word文档读写和模板渲染
"""
from docx import Document
from docxtpl import DocxTemplate
from typing import Dict, Any, List, Optional
from pathlib import Path
import re
from .base_tool import RPAToolBase


class WordTool(RPAToolBase):
    """Word文档处理工具"""
    
    def __init__(self):
        super().__init__()
        self.description = "Word文档读写和模板处理工具"
        self.current_doc = None
    
    def execute(self, **kwargs) -> Dict[str, Any]:
        """通用执行接口"""
        return {"status": "success", "message": "请使用具体方法"}
    
    # ========== 文档读写 ==========
    
    def open_document(self, file_path: str) -> Dict[str, Any]:
        """打开Word文档"""
        try:
            self.current_doc = Document(file_path)
            
            return {
                "status": "success",
                "path": file_path,
                "paragraphs": len(self.current_doc.paragraphs),
                "tables": len(self.current_doc.tables),
                "message": f"打开文档: {file_path}"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def save_document(self, file_path: str) -> Dict[str, Any]:
        """保存Word文档"""
        try:
            if self.current_doc is None:
                return {"status": "error", "message": "没有打开的文档"}
            
            self.current_doc.save(file_path)
            
            return {
                "status": "success",
                "path": file_path,
                "message": f"文档已保存: {file_path}"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    # ========== 内容提取 ==========
    
    def extract_text(self, file_path: Optional[str] = None) -> Dict[str, Any]:
        """
        提取文档全部文本
        
        Args:
            file_path: 文件路径（如果为None则使用当前打开的文档）
        """
        try:
            if file_path:
                doc = Document(file_path)
            elif self.current_doc:
                doc = self.current_doc
            else:
                return {"status": "error", "message": "请先打开文档或指定文件路径"}
            
            text = '\n'.join([p.text for p in doc.paragraphs])
            
            return {
                "status": "success",
                "text": text,
                "length": len(text),
                "paragraphs": len(doc.paragraphs),
                "message": f"提取文本完成: {len(text)} 字符"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def extract_info_by_regex(self, file_path: str, patterns: Dict[str, str]) -> Dict[str, Any]:
        """
        使用正则表达式提取信息（基于现有word_process项目）
        
        Args:
            file_path: Word文件路径
            patterns: 字段名到正则表达式的映射
                例如: {"编号": r"GD-\\d+", "姓名": r"姓名[:：]\\s*([^\\s]+)"}
        """
        try:
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
            
            extracted = {}
            for field_name, pattern in patterns.items():
                match = re.search(pattern, text, re.S)
                if match:
                    # 如果有捕获组，使用第一个捕获组，否则使用整个匹配
                    if match.lastindex and match.lastindex >= 1:
                        extracted[field_name] = match.group(1).strip()
                    else:
                        extracted[field_name] = match.group(0).strip()
                else:
                    extracted[field_name] = "未提取"
            
            return {
                "status": "success",
                "data": extracted,
                "message": f"提取完成: {len(extracted)} 个字段"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    # ========== 模板渲染 ==========
    
    def render_template(self, template_path: str, data: Dict[str, Any],
                       output_path: str) -> Dict[str, Any]:
        """
        渲染Word模板
        
        Args:
            template_path: 模板文件路径
            data: 要填充的数据字典
            output_path: 输出文件路径
        """
        try:
            tpl = DocxTemplate(template_path)
            tpl.render(data)
            tpl.save(output_path)
            
            return {
                "status": "success",
                "template": template_path,
                "output": output_path,
                "fields": list(data.keys()),
                "message": f"模板渲染完成: {output_path}"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    # ========== 内容操作 ==========
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> Dict[str, Any]:
        """添加段落"""
        try:
            if self.current_doc is None:
                return {"status": "error", "message": "请先打开文档"}
            
            self.current_doc.add_paragraph(text, style=style)
            
            return {
                "status": "success",
                "text": text,
                "message": "段落已添加"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def add_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """添加标题"""
        try:
            if self.current_doc is None:
                return {"status": "error", "message": "请先打开文档"}
            
            self.current_doc.add_heading(text, level=level)
            
            return {
                "status": "success",
                "text": text,
                "level": level,
                "message": f"添加{level}级标题"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def add_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> Dict[str, Any]:
        """
        添加表格
        
        Args:
            rows: 行数
            cols: 列数
            data: 可选，表格数据（二维列表）
        """
        try:
            if self.current_doc is None:
                return {"status": "error", "message": "请先打开文档"}
            
            table = self.current_doc.add_table(rows=rows, cols=cols)
            
            # 填充数据
            if data:
                for i, row_data in enumerate(data):
                    if i >= rows:
                        break
                    for j, cell_data in enumerate(row_data):
                        if j >= cols:
                            break
                        table.rows[i].cells[j].text = str(cell_data)
            
            return {
                "status": "success",
                "rows": rows,
                "cols": cols,
                "message": f"添加表格: {rows}x{cols}"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    # ========== 高级功能 ==========
    
    def batch_generate_from_template(self, template_path: str, 
                                     data_list: List[Dict[str, Any]],
                                     output_dir: str,
                                     filename_pattern: str = "{index}.docx") -> Dict[str, Any]:
        """
        批量生成文档（基于模板）
        
        Args:
            template_path: 模板路径
            data_list: 数据列表（每个元素是一个字典）
            output_dir: 输出目录
            filename_pattern: 文件名模式（可使用{index}和数据字段）
        """
        try:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            generated_files = []
            
            for idx, data in enumerate(data_list, start=1):
                # 生成文件名
                filename = filename_pattern.format(index=idx, **data)
                output_file = output_path / filename
                
                # 渲染模板
                tpl = DocxTemplate(template_path)
                tpl.render(data)
                tpl.save(str(output_file))
                
                generated_files.append(str(output_file))
            
            return {
                "status": "success",
                "count": len(generated_files),
                "files": generated_files,
                "message": f"批量生成完成: {len(generated_files)} 个文件"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
